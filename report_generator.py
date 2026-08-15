"""
report_generator.py

Exports the analysis (evidence matrix, gap analysis, personal statement,
optimisation suggestions) as a DOCX or PDF report.

Structured in two layers, same pattern as app.py:

1. build_report_content() — a PURE function. Takes the analysis results
   already produced by earlier modules and assembles them into a
   format-agnostic ReportContent object (title/sections/tables). No file
   I/O, no formatting library calls — fully unit-testable.

2. write_docx_report() / write_pdf_report() — take a ReportContent and
   render it using python-docx / reportlab respectively. Both consume
   the SAME content model, so the DOCX and PDF outputs never drift out
   of sync with each other — there's one source of truth for what goes
   in the report, and two renderers for how it looks on the page.

Nothing in this module invents content. Every section is built directly
from data already produced (and already tested) by cv_parser, evidence_
mapper, gap_analyser, personal_statement, statement_quality, and
optimiser — this module's only job is presentation.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from analysis_bundle import AnalysisBundle
from personal_statement import PersonalStatementResult
from statement_quality import StatementQualityResult
from optimiser import OptimisationResult
import reference_data


# ---------------------------------------------------------------------------
# Format-agnostic content model
# ---------------------------------------------------------------------------

@dataclass
class ReportTable:
    headers: list[str]
    rows: list[list[str]]


@dataclass
class ReportSection:
    heading: str
    paragraphs: list[str] = field(default_factory=list)
    table: ReportTable | None = None


@dataclass
class ReportContent:
    title: str
    subtitle: str
    sections: list[ReportSection] = field(default_factory=list)


def build_report_content(
    bundle: AnalysisBundle,
    target_role: str = "",
    stmt_result: PersonalStatementResult | None = None,
    quality_result: StatementQualityResult | None = None,
    optimisation_result: OptimisationResult | None = None,
    country: str = "",
    industry: str = "",
) -> ReportContent:
    subtitle = f"Prepared for: {target_role}" if target_role else "CV & Job Match Analysis"

    sections = [
        _build_cv_overview_section(bundle),
        _build_jd_overview_section(bundle),
        _build_evidence_matrix_section(bundle),
        _build_gap_summary_section(bundle),
    ]

    if stmt_result is not None:
        sections.append(_build_statement_section(stmt_result, quality_result))

    if optimisation_result is not None:
        sections.append(_build_optimisation_section(optimisation_result))

    context_section = _build_reference_context_section(country, industry)
    if context_section is not None:
        sections.append(context_section)

    return ReportContent(
        title="CV Analyzer & Job Match Report",
        subtitle=subtitle,
        sections=sections,
    )


def _build_cv_overview_section(bundle: AnalysisBundle) -> ReportSection:
    cv_result = bundle.cv_result
    if not cv_result.is_valid:
        return ReportSection(
            heading="CV Overview",
            paragraphs=[f"The CV could not be read: {cv_result.error}"],
        )

    detected = [
        name.capitalize()
        for name in ("summary", "skills", "experience", "education", "certifications", "projects")
        if cv_result.section(name)
    ]
    return ReportSection(
        heading="CV Overview",
        paragraphs=[
            f"CV format: {cv_result.source_format.upper()}",
            f"Sections detected: {', '.join(detected) if detected else 'None detected'}",
        ],
    )


def _build_jd_overview_section(bundle: AnalysisBundle) -> ReportSection:
    sj = bundle.similar_jobs_result
    paragraphs = [
        f"Target job description: {'parsed successfully' if sj.target_jd.is_valid else 'could not be parsed'}",
        f"Similar job descriptions analysed: {sj.similar_jds_parsed} (of {sj.similar_jds_requested} provided)",
    ]
    if sj.warning:
        paragraphs.append(f"Note: {sj.warning}")
    return ReportSection(heading="Job Description Analysis", paragraphs=paragraphs)


def _build_evidence_matrix_section(bundle: AnalysisBundle) -> ReportSection:
    rows = sorted(bundle.evidence_matrix.rows, key=lambda r: r.jd_frequency, reverse=True)
    table = ReportTable(
        headers=["Requirement", "JD Frequency", "Target JD", "CV Evidence", "Source"],
        rows=[
            [
                row.requirement,
                row.jd_frequency,
                row.target_jd_category or "-",
                row.cv_evidence_label,
                row.evidence_source or "-",
            ]
            for row in rows
        ],
    )
    return ReportSection(
        heading="Evidence Matrix",
        paragraphs=[
            "For every requirement identified across the job description(s), this table shows "
            "how often it appeared, how the target role framed it, and whether your CV provides "
            "genuine evidence for it."
        ],
        table=table,
    )


def _build_gap_summary_section(bundle: AnalysisBundle) -> ReportSection:
    gap_result = bundle.gap_result
    paragraphs = []

    if gap_result.high_priority_gaps:
        paragraphs.append("High priority gaps:")
        for entry in gap_result.high_priority_gaps:
            paragraphs.append(f"  - {entry.requirement}: {entry.message}")
    else:
        paragraphs.append("High priority gaps: none — every required item has at least some evidence.")

    if gap_result.needs_stronger_evidence:
        paragraphs.append("Needs stronger evidence:")
        for entry in gap_result.needs_stronger_evidence:
            paragraphs.append(f"  - {entry.requirement}: {entry.message}")

    if gap_result.transferable:
        paragraphs.append("Transferable evidence:")
        for entry in gap_result.transferable:
            paragraphs.append(f"  - {entry.requirement}: {entry.message}")

    paragraphs.append(
        f"Totals: {len(gap_result.strong_matches)} strong match(es), "
        f"{len(gap_result.transferable)} transferable, "
        f"{len(gap_result.needs_stronger_evidence)} needing stronger evidence, "
        f"{len(gap_result.genuine_gaps)} genuine gap(s)."
    )

    return ReportSection(heading="Gap Analysis Summary", paragraphs=paragraphs)


def _build_statement_section(
    stmt_result: PersonalStatementResult, quality_result: StatementQualityResult | None
) -> ReportSection:
    paragraphs = list(stmt_result.statement_text.split("\n\n"))
    if stmt_result.limitation_note:
        paragraphs.append(f"Note: {stmt_result.limitation_note}")

    if quality_result is not None:
        paragraphs.append(
            f"Quality scores — Coverage: {quality_result.coverage_score * 100:.0f}%, "
            f"Evidence: {quality_result.evidence_score * 100:.0f}%, "
            f"Keyword relevance: {quality_result.keyword_relevance_score * 100:.0f}%, "
            f"Natural language: {quality_result.natural_language_score * 100:.0f}%."
        )
        paragraphs.append(quality_result.summary_note)

    return ReportSection(
        heading=f"Personal Statement ({stmt_result.tier.capitalize()})",
        paragraphs=paragraphs,
    )


def _build_optimisation_section(optimisation_result: OptimisationResult) -> ReportSection:
    paragraphs = [optimisation_result.disclaimer]

    if optimisation_result.suggestions:
        paragraphs.append("Suggested wording changes:")
        for s in optimisation_result.suggestions:
            paragraphs.append(f"  [{s.requirement}] Original: {s.original_wording}")
            paragraphs.append(f"  [{s.requirement}] Proposed: {s.proposed_wording}")
            paragraphs.append(f"  [{s.requirement}] Reason: {s.reason}")

    if optimisation_result.recommendations:
        paragraphs.append("Recommendations:")
        for r in optimisation_result.recommendations:
            paragraphs.append(f"  - {r.requirement}: {r.advice}")

    if optimisation_result.unaddressed_gaps:
        paragraphs.append(
            "Not addressed (no evidence in the CV — do not add without genuine experience): "
            + ", ".join(optimisation_result.unaddressed_gaps)
        )

    return ReportSection(heading="CV Optimisation Suggestions", paragraphs=paragraphs)


def _build_reference_context_section(country: str, industry: str) -> ReportSection | None:
    """
    Purely supplementary — never influences the evidence matrix, gap
    analysis, or personal statement above. Returns None (no section at
    all) if neither country nor industry was provided, or if neither
    lookup found a match, rather than showing an empty/awkward section.
    """
    country_ref = reference_data.load_country_reference(country) if country else None
    industry_ref = reference_data.load_industry_reference(industry) if industry else None

    if country_ref is None and industry_ref is None:
        return None

    paragraphs: list[str] = [
        "General background only — the job description above remains the "
        "primary source of truth for this application."
    ]

    for ref in (country_ref, industry_ref):
        if ref is None:
            continue
        paragraphs.append(f"{ref.name}:")
        for convention in ref.key_conventions:
            paragraphs.append(f"  - {convention}")
        if ref.regulatory_considerations:
            paragraphs.append(f"{ref.name} — regulatory considerations:")
            for item in ref.regulatory_considerations:
                paragraphs.append(f"  - {item}")
        if ref.terminology_notes:
            paragraphs.append(f"{ref.name} — terminology notes:")
            for item in ref.terminology_notes:
                paragraphs.append(f"  - {item}")
        paragraphs.append(ref.source_note)

    return ReportSection(heading="Country & Industry Context", paragraphs=paragraphs)


# ---------------------------------------------------------------------------
# DOCX rendering
# ---------------------------------------------------------------------------

def write_docx_report(content: ReportContent, output_path: str | Path) -> Path:
    import docx

    document = docx.Document()

    document.add_heading(content.title, level=0)
    subtitle_paragraph = document.add_paragraph(content.subtitle)
    subtitle_paragraph.runs[0].italic = True

    for section in content.sections:
        document.add_heading(section.heading, level=1)
        for paragraph_text in section.paragraphs:
            document.add_paragraph(paragraph_text)

        if section.table:
            _add_docx_table(document, section.table)

    output_path = Path(output_path)
    document.save(str(output_path))
    return output_path


def _add_docx_table(document, table: ReportTable) -> None:
    docx_table = document.add_table(rows=1, cols=len(table.headers))
    docx_table.style = "Light Grid Accent 1"

    header_cells = docx_table.rows[0].cells
    for i, header in enumerate(table.headers):
        header_cells[i].text = header
        for run in header_cells[i].paragraphs[0].runs:
            run.bold = True

    for row_data in table.rows:
        row_cells = docx_table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)


# ---------------------------------------------------------------------------
# PDF rendering
# ---------------------------------------------------------------------------

def write_pdf_report(content: ReportContent, output_path: str | Path) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    output_path = Path(output_path)
    styles = getSampleStyleSheet()
    subtitle_style = ParagraphStyle(
        "Subtitle", parent=styles["Normal"], fontName="Helvetica-Oblique", spaceAfter=12
    )

    doc = SimpleDocTemplate(
        str(output_path), pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm, leftMargin=2 * cm, rightMargin=2 * cm,
    )
    story = [
        Paragraph(content.title, styles["Title"]),
        Paragraph(content.subtitle, subtitle_style),
        Spacer(1, 12),
    ]

    for section in content.sections:
        story.append(Paragraph(section.heading, styles["Heading2"]))
        for paragraph_text in section.paragraphs:
            # Reportlab's Paragraph treats text as (limited) HTML — escape
            # the handful of characters that would otherwise be
            # misinterpreted as markup, so genuine CV/JD text with & or <
            # in it renders correctly instead of breaking layout.
            safe_text = (
                paragraph_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            )
            story.append(Paragraph(safe_text, styles["Normal"]))
            story.append(Spacer(1, 4))

        if section.table:
            story.append(Spacer(1, 8))
            story.append(_build_pdf_table(section.table))
            story.append(Spacer(1, 12))

    doc.build(story)
    return output_path


def _build_pdf_table(table: ReportTable):
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle

    data = [table.headers] + [[str(v) for v in row] for row in table.rows]
    pdf_table = Table(data, repeatRows=1)
    pdf_table.setStyle(
        TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E5266")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ])
    )
    return pdf_table
