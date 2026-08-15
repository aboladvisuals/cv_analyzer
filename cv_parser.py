"""
cv_parser.py

Responsible for turning a CV — as a PDF, DOCX, TXT file, or pasted plain
text — into clean text, split into recognisable sections (Summary, Skills,
Experience, Education, Certifications, Projects).

This module does ONE job: get reliable text out of whatever the candidate
gave us, and make a first-pass guess at CV structure. It does NOT score,
match, or judge the CV — that's later modules' job (evidence_mapper.py,
gap_analyser.py).

Design choices worth knowing about:
  - Extraction errors never crash the app. A malformed or empty file
    returns a CVParseResult with is_valid=False and a human-readable
    error, so the caller (app.py, or later the Streamlit UI) can show a
    clear message instead of a stack trace.
  - Section detection is rule-based (regex against common header wording).
    It won't be perfect on every CV layout — that's expected and fine for
    this phase. It's a best-effort structure, not a guarantee.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
import re

import config


# ---------------------------------------------------------------------------
# Data shapes
# ---------------------------------------------------------------------------

@dataclass
class CVParseResult:
    is_valid: bool
    raw_text: str = ""
    sections: dict[str, str] = field(default_factory=dict)
    source_format: str = ""       # "pdf" | "docx" | "txt" | "pasted"
    error: str = ""               # populated only when is_valid is False

    def section(self, name: str) -> str:
        """Convenience getter — returns '' if the section wasn't detected."""
        return self.sections.get(name, "")


# ---------------------------------------------------------------------------
# Section headers we look for.
# Each key is the canonical section name; each value is a list of header
# phrases (case-insensitive) that commonly introduce that section in CVs.
# ---------------------------------------------------------------------------

SECTION_HEADERS: dict[str, list[str]] = {
    "summary": ["summary", "profile", "personal profile", "professional summary", "about me"],
    "skills": ["skills", "key skills", "technical skills", "core competencies", "competencies"],
    "experience": ["experience", "work experience", "employment history", "professional experience", "career history"],
    "education": ["education", "academic background", "qualifications"],
    "certifications": ["certifications", "certificates", "professional certifications", "licences", "licenses"],
    "projects": ["projects", "key projects", "personal projects", "portfolio"],
}


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def parse_cv(source: str | Path) -> CVParseResult:
    """
    Parse a CV from either a file path (str or Path, .pdf/.docx/.txt) or
    directly from a string of pasted CV text.

    We treat `source` as pasted text if it isn't a path that exists on
    disk with a supported extension — this keeps the function usable both
    from a CLI/tests context and from a future upload widget.
    """
    if isinstance(source, Path) or _looks_like_file_path(source):
        return _parse_file(Path(source))
    return _parse_pasted_text(str(source))


# ---------------------------------------------------------------------------
# Internal: routing
# ---------------------------------------------------------------------------

def _looks_like_file_path(source: str) -> bool:
    """
    Heuristic: pasted CV text is normally many lines / no valid extension.
    A genuine file path is short, has a supported suffix, and exists.
    """
    if "\n" in source:
        return False
    suffix = Path(source).suffix.lower()
    return suffix in config.SUPPORTED_CV_FORMATS and Path(source).exists()


def _parse_file(path: Path) -> CVParseResult:
    if not path.exists():
        return CVParseResult(is_valid=False, error=f"File not found: {path}")

    suffix = path.suffix.lower()
    if suffix not in config.SUPPORTED_CV_FORMATS:
        supported = ", ".join(config.SUPPORTED_CV_FORMATS)
        return CVParseResult(
            is_valid=False,
            error=f"Unsupported file type '{suffix}'. Supported types: {supported}",
        )

    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb > config.MAX_CV_FILE_SIZE_MB:
        return CVParseResult(
            is_valid=False,
            error=f"File is {size_mb:.1f}MB, which exceeds the {config.MAX_CV_FILE_SIZE_MB}MB limit.",
        )

    try:
        if suffix == ".pdf":
            raw_text = _extract_pdf_text(path)
            source_format = "pdf"
        elif suffix == ".docx":
            raw_text = _extract_docx_text(path)
            source_format = "docx"
        else:  # .txt
            raw_text = _extract_txt_text(path)
            source_format = "txt"
    except Exception as exc:  # noqa: BLE001 — we deliberately want to catch
        # anything the underlying libraries might raise (corrupt file,
        # encoding issues, etc.) and turn it into a clean, user-facing
        # error rather than letting the app crash.
        return CVParseResult(
            is_valid=False,
            source_format=suffix.lstrip("."),
            error=f"Could not read this {suffix} file — it may be corrupted or password-protected. ({exc})",
        )

    return _finish_parsing(raw_text, source_format)


def _parse_pasted_text(text: str) -> CVParseResult:
    return _finish_parsing(text, "pasted")


def _finish_parsing(raw_text: str, source_format: str) -> CVParseResult:
    cleaned = _clean_text(raw_text)

    if not cleaned:
        return CVParseResult(
            is_valid=False,
            source_format=source_format,
            error="No readable text was found. The file may be empty, scanned as an image, or blank.",
        )

    sections = _detect_sections(cleaned)

    return CVParseResult(
        is_valid=True,
        raw_text=cleaned,
        sections=sections,
        source_format=source_format,
    )


# ---------------------------------------------------------------------------
# Internal: format-specific extraction
# ---------------------------------------------------------------------------

def _extract_pdf_text(path: Path) -> str:
    import pdfplumber

    text_parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_docx_text(path: Path) -> str:
    import docx  # python-docx

    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs]

    # Also pull text out of tables — some CVs use a table-based layout for
    # skills or a two-column design, and skipping tables would silently
    # drop that content.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def _extract_txt_text(path: Path) -> str:
    # Try UTF-8 first, fall back to latin-1 for older/Windows-saved files
    # rather than letting an encoding error take down the whole parse.
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


# ---------------------------------------------------------------------------
# Internal: cleaning + section detection
# ---------------------------------------------------------------------------

def _clean_text(text: str) -> str:
    if not text:
        return ""
    # Normalise line endings, collapse excessive blank lines, strip
    # trailing whitespace per line — but keep line breaks, since they
    # carry structural meaning we rely on for section detection.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    cleaned_lines: list[str] = []
    blank_run = 0
    for line in lines:
        if line.strip() == "":
            blank_run += 1
            if blank_run > 1:
                continue
        else:
            blank_run = 0
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines).strip()


def _detect_sections(text: str) -> dict[str, str]:
    """
    Walk the CV line by line. Whenever a line matches one of our known
    section headers, start capturing subsequent lines into that section
    until the next recognised header (or end of document).

    This is intentionally simple: exact-ish header matching rather than
    layout/font analysis (which we don't have access to from plain text
    anyway, especially for PDFs). Good enough to bootstrap the rest of
    the pipeline; can be improved later without changing its interface.
    """
    # Build one lookup: lowercased header phrase -> canonical section name.
    # e.g. "key skills" -> "skills", "employment history" -> "experience"
    phrase_to_section: dict[str, str] = {}
    for canonical_name, phrases in SECTION_HEADERS.items():
        for phrase in phrases:
            phrase_to_section[phrase.lower()] = canonical_name

    # A line counts as a "header line" if, once we strip whitespace and a
    # trailing colon, it matches one of our known phrases exactly. This
    # deliberately does NOT match headers buried mid-sentence — CVs almost
    # always put section headers on their own line.
    lines = text.split("\n")
    sections: dict[str, list[str]] = {name: [] for name in SECTION_HEADERS}
    current_section: str | None = None

    for line in lines:
        stripped_lower = line.strip().rstrip(":").lower()
        if stripped_lower in phrase_to_section:
            current_section = phrase_to_section[stripped_lower]
            continue
        if current_section:
            sections[current_section].append(line)

    return {
        name: "\n".join(content_lines).strip()
        for name, content_lines in sections.items()
        if content_lines
    }
